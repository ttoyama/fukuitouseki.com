#!/usr/bin/env python3
"""
Docx to Markdown converter script
Converts Word documents to Markdown format with proper formatting preservation
"""

import os
import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH

def extract_text_with_formatting(paragraph):
    """Extract text from paragraph while preserving formatting"""
    text = ""
    for run in paragraph.runs:
        run_text = run.text
        
        # Apply formatting
        if run.bold and run.italic:
            run_text = f"***{run_text}***"
        elif run.bold:
            run_text = f"**{run_text}**"
        elif run.italic:
            run_text = f"*{run_text}*"
        
        text += run_text
    
    return text

def process_paragraph(paragraph):
    """Process a paragraph and convert to appropriate markdown"""
    text = extract_text_with_formatting(paragraph)
    
    # Skip empty paragraphs
    if not text.strip():
        return ""
    
    # Convert checkboxes
    text = re.sub(r'☐|□|◇|○', '- [ ]', text)
    text = re.sub(r'☑|■|◆|●', '- [x]', text)
    
    # Handle different paragraph styles
    style_name = paragraph.style.name.lower() if paragraph.style else ""
    
    # Headers
    if 'heading' in style_name or 'title' in style_name:
        if '1' in style_name or 'title' in style_name:
            return f"# {text}\n"
        elif '2' in style_name:
            return f"## {text}\n"
        elif '3' in style_name:
            return f"### {text}\n"
        elif '4' in style_name:
            return f"#### {text}\n"
        else:
            return f"## {text}\n"
    
    # Check for manual formatting that looks like headers
    if paragraph.runs:
        first_run = paragraph.runs[0]
        if first_run.bold and len(text.strip()) < 100:
            # Likely a header
            return f"## {text}\n"
    
    # Regular paragraph
    return f"{text}\n"

def process_table(table):
    """Convert docx table to markdown table"""
    if not table.rows:
        return ""
    
    markdown_table = []
    
    # Process header row
    header_row = table.rows[0]
    headers = []
    for cell in header_row.cells:
        cell_text = ""
        for paragraph in cell.paragraphs:
            cell_text += extract_text_with_formatting(paragraph) + " "
        headers.append(cell_text.strip())
    
    if headers:
        markdown_table.append("| " + " | ".join(headers) + " |")
        markdown_table.append("| " + " | ".join(["---"] * len(headers)) + " |")
    
    # Process data rows
    for row in table.rows[1:]:
        row_data = []
        for cell in row.cells:
            cell_text = ""
            for paragraph in cell.paragraphs:
                cell_text += extract_text_with_formatting(paragraph) + " "
            # Convert checkboxes in table cells
            cell_text = re.sub(r'☐|□|◇|○', '- [ ]', cell_text)
            cell_text = re.sub(r'☑|■|◆|●', '- [x]', cell_text)
            row_data.append(cell_text.strip())
        
        if row_data:
            markdown_table.append("| " + " | ".join(row_data) + " |")
    
    return "\n".join(markdown_table) + "\n\n"

def docx_to_markdown(docx_path):
    """Convert a docx file to markdown"""
    try:
        doc = Document(docx_path)
        markdown_content = []
        
        # Process document elements
        for element in doc.element.body:
            if element.tag.endswith('p'):  # Paragraph
                for paragraph in doc.paragraphs:
                    if paragraph._element == element:
                        md_text = process_paragraph(paragraph)
                        if md_text:
                            markdown_content.append(md_text)
                        break
            elif element.tag.endswith('tbl'):  # Table
                for table in doc.tables:
                    if table._element == element:
                        md_table = process_table(table)
                        if md_table:
                            markdown_content.append(md_table)
                        break
        
        return "".join(markdown_content)
    
    except Exception as e:
        print(f"Error processing {docx_path}: {str(e)}")
        return None

def convert_file(docx_path, output_dir=None):
    """Convert a single docx file to markdown"""
    docx_path = Path(docx_path)
    
    if not docx_path.exists():
        print(f"Error: File {docx_path} does not exist")
        return False
    
    if docx_path.suffix.lower() != '.docx':
        print(f"Error: {docx_path} is not a docx file")
        return False
    
    # Generate output filename
    if output_dir:
        output_dir = Path(output_dir)
        output_dir.mkdir(exist_ok=True)
        output_path = output_dir / f"{docx_path.stem}.md"
    else:
        output_path = docx_path.parent / f"{docx_path.stem}.md"
    
    print(f"Converting {docx_path.name} to {output_path.name}...")
    
    # Convert to markdown
    markdown_content = docx_to_markdown(docx_path)
    
    if markdown_content is None:
        return False
    
    # Write output file
    try:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        print(f"Successfully converted to {output_path}")
        return True
    except Exception as e:
        print(f"Error writing to {output_path}: {str(e)}")
        return False

def main():
    """Main function"""
    if len(sys.argv) < 2:
        print("Usage: python docx_to_markdown.py <docx_file_or_directory> [output_directory]")
        print("Examples:")
        print("  python docx_to_markdown.py document.docx")
        print("  python docx_to_markdown.py docs/資料/ converted/")
        print("  python docx_to_markdown.py document.docx output/")
        sys.exit(1)
    
    input_path = Path(sys.argv[1])
    output_dir = sys.argv[2] if len(sys.argv) > 2 else None
    
    if input_path.is_file():
        # Convert single file
        success = convert_file(input_path, output_dir)
        if not success:
            sys.exit(1)
    elif input_path.is_dir():
        # Convert all docx files in directory
        docx_files = list(input_path.glob("*.docx"))
        if not docx_files:
            print(f"No docx files found in {input_path}")
            sys.exit(1)
        
        success_count = 0
        for docx_file in docx_files:
            if convert_file(docx_file, output_dir):
                success_count += 1
        
        print(f"\nConversion complete: {success_count}/{len(docx_files)} files converted successfully")
    else:
        print(f"Error: {input_path} is not a valid file or directory")
        sys.exit(1)

if __name__ == "__main__":
    main()